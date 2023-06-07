from html.parser import HTMLParser
from docx import Document
import io
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
import base64
import re
from bs4 import BeautifulSoup
from docx.oxml import OxmlElement as OE
from docx.oxml.ns import qn
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX

styling_tags = ["strong", "em", "u", "span", "a"]
custom_vars_pat = r"\[\%[\w\s]+\%\]"


def isBase64(s):
    try:
        b64 = base64.b64encode(base64.b64decode(s))
        return len(b64) > 0 and b64.decode("utf-8") == s
    except Exception:
        return False


def highlight_word(para, word):
    para_text = para.text
    para.text = ""
    para_lines = para_text.split(word)
    for para_line in para_lines[0 : len(para_lines) - 1]:
        para.add_run(para_line)
        para.add_run(word).font.highlight_color = WD_COLOR_INDEX.YELLOW
    para.add_run(para_lines[len(para_lines) - 1])
    return para


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(
        docx.oxml.shared.qn("r:id"),
        r_id,
    )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement("w:r")
    rPr = docx.oxml.shared.OxmlElement("w:rPr")

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


class DocHTMLParser(HTMLParser):
    def set_table_tags(self):
        self.cur_col = -1
        self.cur_row = -1
        self.cur_cell_type = ""
        self.total_rows = 0
        self.total_cols = 0

    def set_prev_start_tag(self):
        self.prev_start_tag = ""
        self.last_start_tag = ""

    def set_fig_count(self):
        self.fig_count = 0

    def set_table_count(self):
        self.table_count = 0

    def handle_starttag(self, tag, attrs):
        self.last_end_tag = ""
        if tag not in styling_tags:
            self.is_data = False
            self.prev_start_tag = self.last_start_tag
        self.last_start_tag = tag
        self.last_attrs = attrs
        if tag == "table":
            self.table_count += 1
            for attr in attrs:
                attr_name, attr_value = attr
                if attr_name == "rows":
                    self.total_rows = int(attr_value)
                if attr_name == "columns":
                    self.total_cols = int(attr_value)
            self.word_doc.add_table(rows=self.total_rows, cols=self.total_cols)
        if tag == "tr":
            self.cur_row += 1
        if tag == "td" or tag == "th":
            self.cur_cell_type = tag
            self.cur_col += 1
        if len(re.findall("[uo]l", tag)) > 0:
            self.cur_list_type = tag
        if tag == "img":
            img_src = ""
            inch_const = 0.010417
            img_width = 500 * inch_const
            img_height = 200 * inch_const
            self.fig_count = self.fig_count + 1
            for attr in attrs:
                attr_name, attr_value = attr
                if attr_name == "src":
                    img_src = attr_value.split("base64,")[1]
                if attr_name == "width":
                    img_width = float(attr_value) * inch_const
                if attr_name == "height":
                    img_height = float(attr_value) * inch_const
            img_bytes = base64.b64decode(img_src)
            image = io.BytesIO(img_bytes)
            self.word_doc.add_picture(
                image, width=Inches(img_width), height=Inches(img_height)
            )
            last_paragraph = self.word_doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.word_doc.add_paragraph(
                f"Figure {self.fig_count}: Figure Title Here", style="186-Figure-Title"
            )

    def handle_endtag(self, tag):
        self.last_end_tag = tag
        if tag == "table":
            self.total_rows = 0
            self.total_cols = 0
            self.cur_row = -1
            self.word_doc.add_paragraph(
                f"Table {self.table_count}: Table Title Here", style="186-Figure-Title"
            )
        if tag == "tr":
            self.cur_col = -1
        if tag == "td" or tag == "th":
            self.cur_cell_type = ""
        if len(re.findall("[uo]l", tag)) > 0:
            self.cur_list_type = ""
        last_paragraph = self.word_doc.paragraphs[-1]
        if len(re.findall("li", self.prev_start_tag)) > 0:
            if self.cur_list_type == "ul":
                last_paragraph.style = "031-Bulleted-List1"
            elif self.cur_list_type == "ol":
                last_paragraph.style = "040-Numbered-List"

    def handle_data(self, data):
        if len(self.cur_cell_type) > 0 and self.total_rows > 0 and self.total_cols > 0:
            cur_table = self.word_doc.tables[-1]
            if self.cur_cell_type == "th":
                hdr_cells = cur_table.rows[0].cells
                try:
                    hdr_cells[self.cur_col].add_paragraph(data, "100-Paragraph")
                except Exception as e:
                    pass
            elif self.cur_cell_type == "td":
                if len(cur_table.rows) > self.cur_row:
                    pass
                else:
                    cur_table.add_row().cells
                row_cells = cur_table.rows[self.cur_row].cells
                row_cells[self.cur_col].add_paragraph(data, "100-Paragraph")
        elif len(re.findall("Table [0-9]*:", data)) > 0:
            last_paragraph = self.word_doc.paragraphs[-1]
            if self.table_count != 0:
                table_name = data.split(":")
                if len(table_name) > 1:
                    table_name = " ".join(table_name[1:]).strip()
                    last_paragraph.text = f"Table {self.table_count}: {table_name}"
                else:
                    last_paragraph.text = f"Table {self.table_count}: Table Title Here"
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif len(re.findall("Figure [0-9]*:", data)) > 0:
            last_paragraph = self.word_doc.paragraphs[-1]
            if self.fig_count != 0:
                fig_name = data.split(":")
                if len(fig_name) > 1:
                    fig_name = " ".join(fig_name[1:]).strip()
                    last_paragraph.text = f"Figure {self.fig_count}: {fig_name}"
                else:
                    last_paragraph.text = f"Figure {self.fig_count}: Figure Title Here"
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            if self.last_start_tag not in styling_tags or (
                self.last_start_tag in styling_tags and self.is_data == False
            ):
                self.is_data = True
                self.word_doc.add_paragraph()
            last_paragraph = self.word_doc.paragraphs[-1]
            if len(re.findall("h[0-9]", self.last_start_tag)) > 0:
                header_idx = int(self.last_start_tag[1])
                last_paragraph.text = data
                last_paragraph.style = f"Heading {header_idx}"
            elif self.last_start_tag in styling_tags:
                if self.last_end_tag not in styling_tags:
                    tag = self.last_start_tag
                    if tag == "strong":
                        last_paragraph.add_run(data).bold = True
                    elif tag == "em":
                        last_paragraph.add_run(data).italic = True
                    elif tag == "u":
                        last_paragraph.add_run(data).underline = True
                    elif tag == "span" and len(self.last_attrs) > 0:
                        for attr in self.last_attrs:
                            attr_name, attr_value = attr
                            if attr_name == "style":
                                if len(re.findall("^color:", attr_value)) > 0:
                                    # color: rgb(18, 65, 145);
                                    run = last_paragraph.add_run(data)
                                    color = (
                                        attr_value.split(":")[1]
                                        .strip()
                                        .replace(";", "")
                                    )
                                    if color != "rgb(0, 0, 0)":
                                        run.font.color.rgb = RGBColor.from_string(
                                            "005aff"
                                        )
                                if (
                                    len(re.findall("^background-color:", attr_value))
                                    > 0
                                ):
                                    # bg_color = attr_value.split(":")[1].strip().replace(";","")
                                    last_paragraph.add_run(
                                        data
                                    ).font.highlight_color = WD_COLOR_INDEX.YELLOW
                    elif tag == "a":
                        link = ""
                        for attr in self.last_attrs:
                            name, style = attr
                            if name == "href":
                                link = style
                                break
                        add_hyperlink(last_paragraph, data, link)
                    else:
                        last_paragraph.add_run(data)
                else:
                    last_paragraph.add_run(data)
            elif len(re.findall("li", self.last_start_tag)) > 0:
                if self.cur_list_type == "ul":
                    last_paragraph.text = data
                    last_paragraph.style = "031-Bulleted-List1"
                elif self.cur_list_type == "ol":
                    last_paragraph.text = data
                    last_paragraph.style = "040-Numbered-List"
            else:
                last_paragraph.text = data
                last_paragraph.style = "100-Paragraph"

    def set_document(self):
        self.word_doc = Document("doc_assist_be/wrds/Sample File - V3.docx")

    def get_document(self):
        return self.word_doc


def add_list_of_table(run):
    fldChar = OE("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    fldChar.set(qn("w:dirty"), "true")
    instrText = OE("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "TOC \\o \\h \\z \\u \\c"  # "Table" of list of table and "Figure" for list of figure
    fldChar2 = OE("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OE("w:t")
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OE("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar4)


def add_list_of_fig_tab(run, is_fig=False):
    fldChar = OE("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    fldChar.set(qn("w:dirty"), "true")
    instrText = OE("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    if is_fig:
        instrText.text = 'TOC \\z \\u \\c "Figure"'  # "Table" of list of table and "Figure" for list of figure
    else:
        instrText.text = 'TOC \\z \\u \\c "Table"'  # "Table" of list of table and "Figure" for list of figure
    fldChar2 = OE("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OE("w:t")
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OE("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar4)


def fig_tab_caption(paragraph, is_figure=False):
    run = run = paragraph.add_run()
    r = run._r
    fldChar = OE("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    r.append(fldChar)
    instrText = OE("w:instrText")
    if is_figure:
        instrText.text = " SEQ Figure * ARABIC"
    else:
        instrText.text = " SEQ Table * ARABIC"
    r.append(instrText)
    fldChar = OE("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "end")
    r.append(fldChar)


def format_table(cell, para, is_header=False):
    """
    set text color and alignment
    """
    text = para.text
    delete_paragraph(para)
    cur_para = cell.add_paragraph()
    # cur_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cur_para.add_run(text)
    if is_header:
        run.font.color.rgb = RGBColor.from_string("FFFFFF")  # White Color
    return para


def set_table_header_bg_color(cell, bg_color):
    """
    set background shading for Header Rows
    """
    tblCell = cell._tc
    tblCellProperties = tblCell.get_or_add_tcPr()
    clShading = OE("w:shd")
    clShading.set(
        qn("w:fill"), bg_color
    )  # Hex of Dark Blue Shade {R:0x00, G:0x51, B:0x9E}
    tblCellProperties.append(clShading)
    return cell


def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OE("w:tcBorders")
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ("start", "top", "end", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = "w:{}".format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OE(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))


def format_sections(sections):
    html_sections = []
    for section in sections:
        cur_sec = section.split(" ")
        if len(cur_sec) > 0:
            cur_sec_idx = cur_sec[0]
            cur_sec_title = " ".join(cur_sec[1:])
            total_count = list(filter(lambda sec: len(sec) > 0, cur_sec_idx.split(".")))
            if not re.findall(r"Figure [0-9 ]*:", section) and not re.findall(
                r"Table [0-9 ]*:", section
            ):
                header_str = (
                    f"<h{len(total_count)}>{cur_sec_title}</h{len(total_count)}>"
                )
            else:
                header_str = f"<h1>{section}</h1>"
        html_sections.append(header_str)
    return html_sections


def format_html(html_sections, html_arr):
    converted_html_arr = []
    for idx, html in enumerate(html_arr):
        soup = BeautifulSoup(html, "html.parser")
        # Find the table element using its tag name
        if soup.head:
            soup.head.decompose()
        if soup.body:
            tables = soup.body.unwrap().find_all("table")
        else:
            tables = soup.find_all("table")

        for table in tables:
            row_ = table.find_all("tr")
            col_ = row_[0].find_all("td") or row_[0].find_all("th")
            table["rows"] = len(row_)
            table["columns"] = len(col_)

        converted_html_arr.append(html_sections[idx] + str(soup))

    return "".join(converted_html_arr)


def update_custom_vars(
    bbt_data, custom_vars_dict, custom_vars, para, prop, row_idx, ecell_idx
):
    for cvar in custom_vars:
        cvar_pat = re.sub("\[", "\[", cvar)
        cvar_pat = re.sub("\]", "\]", cvar_pat)
        if len(re.findall(cvar_pat, para.text)) > 0:
            if (
                type(custom_vars_dict[cvar]) != str
                and custom_vars_dict[cvar]["crm_bg"] == "crm"
            ):
                option = custom_vars_dict[cvar]["crm_bg_option"]
                if option in bbt_data.keys():
                    para.text = re.sub(cvar_pat, bbt_data[option], para.text)
                else:
                    para = highlight_word(para, cvar)
            elif cvar in prop.keys():
                if isBase64(prop[cvar]):
                    para.text = para.text.replace(cvar, "")
                    cur_run = para.add_run()
                    inch_const = 0.010417
                    img_width = 500 * inch_const
                    img_height = 200 * inch_const
                    img_bytes = base64.b64decode(prop[cvar])
                    image = io.BytesIO(img_bytes)
                    cur_run.add_picture(
                        image,
                        width=Inches(img_width),
                        height=Inches(img_height),
                    )
                elif type(prop[cvar]) == list:
                    if len(prop[cvar][row_idx]) > ecell_idx:
                        para.text = prop[cvar][row_idx][ecell_idx]
                    else:
                        para.text = ""
                elif type(prop[cvar]) == str and len(prop[cvar]) > 0:
                    para.text = para.text = re.sub(cvar_pat, prop[cvar], para.text)
                elif type(prop[cvar]) == str and len(prop[cvar]) == 0:
                    para = highlight_word(para, cvar)
            else:
                para = highlight_word(para, cvar)
    match_ls = re.findall(custom_vars_pat, para.text)
    match_ls = list(filter(lambda cust_var: cust_var not in custom_vars, match_ls))
    if len(match_ls) > 0:
        for cust_var in match_ls:
            para = highlight_word(para, cust_var)
