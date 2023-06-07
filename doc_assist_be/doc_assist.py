import base64
from flask import request, Blueprint, send_file
import requests
from werkzeug.utils import secure_filename
import mammoth
import json
import os
from docx import Document
from docx2pdf import convert
from doc_assist_be.services.extras import (
    bson_to_json,
    create_ver_heirarchy,
    format_data,
    get_bbt_data,
    get_restricted_keywrds_list,
    get_sensitive_sections,
    json_to_bson,
)
from doc_assist_be.services.html_parser import (
    DocHTMLParser,
    add_list_of_fig_tab,
    add_list_of_table,
    delete_paragraph,
    fig_tab_caption,
    format_html,
    format_sections,
    isBase64,
    set_cell_border,
    set_table_header_bg_color,
    format_table,
    update_custom_vars,
)
import re

# import pythoncom
# import os

# import comtypes.client
from natsort import natsorted
from datetime import datetime
import pymongo

client = pymongo.MongoClient(
    "mongodb://docassistdb:iLrkYB7oVaweYFLJdfFTrFwMdLYrMw84rRyNOmOeQhiIKMhSlpahN2nbFNFaF9K93OR0VpEatmwqACDbc4kvxA==@docassistdb.mongo.cosmos.azure.com:10255/?ssl=true&replicaSet=globaldb&retrywrites=false&maxIdleTimeMS=120000&appName=@docassistdb@"
)
db = client["doc-assist-db"]

master_mappings = db["master_mappings"]
download_history_mappings = db["download_history_mappings"]
sections_mappings = db["sections_params_mapping_cache"]
heirarchy_mappings = db["heirarchy_mapping_cache"]
custom_vars_mappings = db["custom_vars_mapping"]
restricted_ls = db["restricted_ls"]
# filepath = r"C:\Users\parjain\Documents\doc-assist\doc-assist\services\converted_doc.docx"  # put file path here
filepath = r"doc_assist_be\services\converted_doc.docx"  # put file path here

custom_vars_pat = r"\[\%[\w\s]+\%\]"

application = Blueprint("doc-assist", __name__)


@application.route("/get-doc-upload-dropdowns", methods=["GET", "POST"])
def get_doc_upload_dropdowns():
    try:
        doc_upload_dropdowns = {
            "bg": [
                {"name": "MN", "value": "MN"},
                {"name": "CNS-CSP", "value": "CNS-CSP"},
                {"name": "NI", "value": "NI"},
                {"name": "CNS Enterprise", "value": "CNS Enterprise"},
            ],
            "functions": [
                {"name": "Operations", "value": "Operations"},
                {"name": "Legal", "value": "Legal"},
            ],
            "access": [
                {"name": "Nokia", "value": "Nokia"},
                {"name": "Limited to user group", "value": "Limited GRP"},
            ],
            "coOwner": [
                {"name": "Nokia", "value": "Nokia"},
                {"name": "Limited to user group", "value": "Limited GRP"},
            ],
            "jsonConfig": [
                {"name": "BG Configurator", "value": "bg_config"},
                {"name": "Manual Generation", "value": "manual_gen"},
                {"name": "Both", "value": "both"},
            ],
        }
        return {"status": "Received!", "doc_upload_dropdowns": doc_upload_dropdowns}
    except Exception as exp:
        return {
            "status": f"Error Occured as {exp}!",
            "doc_upload_dropdowns": {
                "bg": [],
                "functions": [],
                "access": [],
                "coOwner": [],
            },
        }


@application.route("/get-gen-doc-dropdowns", methods=["GET", "POST"])
def get_gen_doc_dropdowns():
    try:
        doc_gen_dropdowns = {
            "bg": [
                {"name": "MN", "value": "MN"},
                {"name": "CNS-CSP", "value": "CNS-CSP"},
                {"name": "NI", "value": "NI"},
                {"name": "CNS Enterprise", "value": "CNS Enterprise"},
            ],
            "functions": [
                {"name": "Operations", "value": "Operations"},
                {"name": "Legal", "value": "Legal"},
            ],
        }
        return {"status": "Received!", "doc_gen_dropdowns": doc_gen_dropdowns}
    except Exception as exp:
        return {
            "status": "Error Occured as {exp}!",
            "doc_gen_dropdowns": {
                "bg": [],
                "functions": [],
            },
        }


@application.route("/version-sort", methods=["GET", "POST"])
def version_sort():
    try:
        prop = request.get_json()
        versions = prop["ver_li"]
        if len(versions) > 0:
            version_li = [
                {
                    "id": version["id"],
                    "ver": version["value"].split(" ")[0],
                    "value": version["value"],
                }
                for version in versions
            ]
            version_li = natsorted(version_li, key=lambda version: version["ver"])
            return {"status": "Received!", "sorted_versions": version_li}
        return {"status": "Received but could not process the request!"}
    except Exception as exp:
        return {"status": f"Error occured as {exp}"}


@application.route("/get_html", methods=["GET", "POST"])
def get_html():
    try:
        file = request.files["file"]
        filename = secure_filename(file.filename)
        file.save(f"doc_assist_be/services/{filename}")
        with open(f"doc_assist_be/services/{filename}", "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)

        return result.value
    except Exception as exp:
        return {"status": f"Error occured as {exp}"}


@application.route("/upload_prop", methods=["GET", "POST"])
def upload_prop():
    try:
        prop = request.get_json()
        domain = list(prop.keys())[0]
        filename = list(prop[domain].keys())[0]
        is_file_exists = True
        if is_file_exists:
            json_db = {}
            master_db = master_mappings.find_one(
                format_data(filename, domain)["filter"]
            )
            if master_db != None:
                json_db = dict(master_db)
                json_db = bson_to_json("master_mappings", json_db, filename, domain)

            if json_db != None:
                # new_param_obj = {}
                if domain in json_db.keys() and filename in json_db[domain].keys():
                    json_db[domain][filename] = prop[domain][filename]
                elif domain in json_db.keys():
                    json_db[domain].update(prop[domain])
                else:
                    json_db[domain] = prop[domain]

                # Sort According to sections
                sec_map = list(
                    map(
                        lambda sec_key: {
                            "mapped_key": sec_key,
                            "section": json_db[domain][filename][sec_key][
                                "section"
                            ].split(" ")[0],
                        },
                        list(json_db[domain][filename].keys()),
                    )
                )
                sorted_sec_map = natsorted(
                    sec_map, key=lambda section: section["section"]
                )
                sorted_file_sections = {}
                for cur_sec in sorted_sec_map:
                    mapped_key = cur_sec["mapped_key"]
                    sorted_file_sections[mapped_key] = json_db[domain][filename][
                        mapped_key
                    ]

                json_db[domain][filename] = sorted_file_sections
                master_bson = json_to_bson(
                    {"mappings": json_db[domain][filename]},
                    format_data(filename, domain)["filter"],
                )
                if bool(master_db):
                    master_mappings.update_one(
                        format_data(filename, domain)["filter"],
                        {"$set": {"mappings": json_db[domain][filename]}},
                    )
                else:
                    master_mappings.insert_many(master_bson)
        else:
            master_bson = json_to_bson(
                {"mappings": prop[domain][filename]},
                format_data(filename, domain)["filter"],
            )
            if bool(master_db):
                master_mappings.update_one(
                    format_data(filename, domain)["filter"],
                    {"$set": {"mappings": prop[domain][filename]}},
                )
            else:
                master_mappings.insert_many(master_bson)
        # Add Custom Vars
        custom_vars_dict = {
            cust_var: cust_val
            for cust_var, cust_val in prop.items()
            if cust_var != domain
        }
        custom_vars_dict_json = {}

        for cust_var in prop.keys():
            if cust_var != domain:
                if type(custom_vars_dict[cust_var]) == str:
                    custom_vars_dict_json[cust_var] = ""
                elif custom_vars_dict[cust_var]["crm_bg"] != "crm":
                    custom_vars_dict_json[cust_var] = custom_vars_dict[cust_var][
                        "crm_bg_option"
                    ].capitalize()

        is_file_exists = True
        if is_file_exists:
            custom_vars_db = {}
            custom_var_db = custom_vars_mappings.find_one(
                format_data(filename, domain)["filter"]
            )
            if custom_var_db != None:
                custom_vars_db = dict(custom_var_db)
                custom_vars_db = bson_to_json(
                    "custom_vars_mappings", custom_vars_db, filename, domain
                )
            if custom_vars_db != None:
                # Update Custom Vars
                if (
                    domain in custom_vars_db.keys()
                    and filename in custom_vars_db[domain].keys()
                ):
                    custom_vars_db[domain][filename] = custom_vars_dict
                elif domain in custom_vars_db.keys():
                    custom_vars_db[domain].update({filename: custom_vars_dict})
                else:
                    custom_vars_db.update({domain: {filename: custom_vars_dict}})
                custom_vars_bson = json_to_bson(
                    {"custom_vars": custom_vars_db[domain][filename]},
                    format_data(filename, domain)["filter"],
                )
                if bool(custom_var_db):
                    custom_vars_mappings.update_one(
                        format_data(filename, domain)["filter"],
                        {"$set": {"custom_vars": custom_vars_db[domain][filename]}},
                    )
                else:
                    custom_vars_mappings.insert_many(custom_vars_bson)
        # Downloadable JSON here
        ascii_enc_bytes_domain = domain.encode("ascii")
        base64_enc_domain = base64.b64encode(ascii_enc_bytes_domain)
        download_json = {
            "file_id": f"{filename}-{base64_enc_domain.decode('ascii')}",
            "key_map": [],
            "is_crm": "",
            "crmid": "",
        }
        download_json.update(custom_vars_dict_json)
        param_ls = prop[domain][filename].keys()
        possible_values = {}
        for param in param_ls:
            params = param.split("_")
            param_dep = f"{params[0]}_{params[1]}"
            value = "_".join(params[2:])
            if param_dep in possible_values.keys():
                possible_values[param_dep].append(value)
            else:
                possible_values[param_dep] = [value]

        for param in possible_values.keys():
            params = param.split("_")
            ini_value = possible_values[f"{params[0]}_{params[1]}"][0]
            if prop[domain][filename][f"{param}_{ini_value}"]["config"] == "manual_gen":
                break
            elif params[1] != "optional" and params[1] != "mandatory":
                download_json["key_map"].append(
                    {
                        "parameter": params[0],
                        "dependency": params[1],
                        "value": "_".join(params[2:]),
                        "Possible Values": "/".join(
                            possible_values[f"{params[0]}_{params[1]}"]
                        ),
                    }
                )

        return {"status": "Received!", "download_json": download_json}
    except Exception as exp:
        return {"status": f"Error occured as {exp}"}


@application.route("/download_doc", methods=["GET", "POST"])
def download_doc():
    try:
        props = request.get_json()
        # props = bytes.decode(request.data, "ascii")
        # props = json.loads(f"""{props}""".replace("'", '"'))
        selected_doc = props["filename"]
        domain = props["domain"]
        key_map = [key.replace("\xa0", " ") for key in props["key_map"]]
        doc_content = []

        if len(selected_doc) > 0:
            is_file_exists = True
            doc_sections = {}
            if is_file_exists:
                json_db = {}
                master_db = master_mappings.find_one(
                    format_data(selected_doc, domain)["filter"]
                )
                if master_db != None:
                    json_db = dict(master_db)
                    json_db = bson_to_json(
                        "master_mappings", json_db, selected_doc, domain
                    )

                if json_db != None and domain in json_db.keys():
                    if selected_doc in list(json_db[domain].keys()):
                        doc_sections = json_db[domain][selected_doc]
            if len(list(doc_sections.keys())) > 0:
                selected_keys = [
                    cur_
                    for cur_ in list(doc_sections.keys())
                    if cur_ in key_map or cur_.split("_")[1] == "mandatory"
                ]
                # Duplicate Sections
                dup_sections = [
                    {
                        "dup_section": doc_sections[section_key]["section"],
                        "content": doc_sections[section_key]["content"],
                    }
                    for section_key in selected_keys
                    if re.search(
                        r"0.[a-z]$", doc_sections[section_key]["section"].split(" ")[0]
                    )
                ]
                if len(dup_sections) > 0:
                    for dup_sec in dup_sections:
                        cur_sec = dup_sec["dup_section"].split(" ")[0]
                        sec_ls = cur_sec.split(".")
                        parent_sec = ".".join(sec_ls[0 : len(sec_ls) - 2])
                        for section_key in selected_keys:
                            section_title = doc_sections[section_key]["section"]
                            section_ = section_title.split(" ")[0]
                            if section_ == parent_sec:
                                doc_sections[section_key]["content"] += dup_sec[
                                    "content"
                                ]
                final_sel_sections = []
                sel_sec_idx = list(
                    map(
                        lambda sec: doc_sections[sec]["section"].split(" ")[0],
                        selected_keys,
                    )
                )
                for section_key in selected_keys:
                    section_title = doc_sections[section_key]["section"]
                    section_idx = section_title.split(" ")[0]
                    # Add Parent Section if Duplicate is selected here!
                    if not re.search(r"0.[a-z]$", section_idx):
                        for parent_idx in range(len(section_idx.split("."))):
                            parent_section = ".".join(
                                section_idx.split(".")[0:parent_idx]
                            )
                            if len(parent_section) > 0:
                                if parent_section in sel_sec_idx:
                                    pass
                                else:
                                    sel_sec_idx.append(parent_section)

                sel_sec_idx = natsorted(sel_sec_idx)
                for section_key in doc_sections.keys():
                    section_title = doc_sections[section_key]["section"]
                    section_idx = section_title.split(" ")[0]
                    if section_idx in sel_sec_idx:
                        doc_content.append(doc_sections[section_key]["content"])
                        final_sel_sections.append(section_title)

        return {
            "status": "Received!",
            "contents": doc_content,
            "sections": final_sel_sections,
        }
    except Exception as exp:
        return {"status": f"Error occured as {exp}"}


@application.route("/get_prop", methods=["GET", "POST"])
def get_prop():
    try:
        prop = request.get_json()
        domain = prop["domain"]
        cur_user = prop["user"]
        props = []
        is_file_exists = True
        if is_file_exists:
            json_db = {}
            master_db = master_mappings.find({"user_domain": domain})
            if master_db != None:
                master_db = list(master_db)
                for file in master_db:
                    json_db[file["user_file"]] = file["mappings"]

                json_db = {domain: json_db}

            if json_db != None:
                if domain in json_db.keys():
                    for doc in json_db[domain].keys():
                        doc_param = list(json_db[domain][doc].keys())[0]
                        access_lvl = json_db[domain][doc][doc_param]["access_level"]
                        allowed_users = list(
                            filter(lambda user: len(user), access_lvl.split(";"))
                        )
                        if cur_user in allowed_users or "Nokia" in allowed_users:
                            doc_name = json_db[domain][doc][doc_param]["doc_name"]
                            doc_desc = json_db[domain][doc][doc_param]["doc_desc"]
                            doc_config = json_db[domain][doc][doc_param]["config"]
                            file_uploaded = doc
                            props.append(
                                {
                                    "doc_name": doc_name,
                                    "doc_desc": doc_desc,
                                    "file_uploaded": file_uploaded,
                                    "config": doc_config,
                                }
                            )

        return {"status": "Success!", "docs": props}
    except Exception as exp:
        return {"status": f"Error occured as {exp}"}


@application.route("/get_sections", methods=["GET", "POST"])
def get_sections():
    try:
        prop = request.get_json()
        selected_doc = prop["selectedDoc"]
        domain = prop["domain"]
        doc_content = []
        # custom_vars = []
        custom_vars_dict = {}
        initial_content = ""
        if len(selected_doc) > 0:
            # Custom Vars
            is_file_exists = True
            if is_file_exists:
                custom_vars_db = custom_vars_mappings.find_one(
                    format_data(selected_doc, domain)["filter"]
                )
                if custom_vars_db != None:
                    custom_vars_payload = dict(custom_vars_db)
                    custom_vars_payload = bson_to_json(
                        "custom_vars_mappings",
                        custom_vars_payload,
                        selected_doc,
                        domain,
                    )
                    if (
                        domain in custom_vars_payload.keys()
                        and selected_doc in custom_vars_payload[domain].keys()
                    ):
                        custom_vars_dict = {
                            cust_var: cust_val
                            for cust_var, cust_val in custom_vars_payload[domain][
                                selected_doc
                            ].items()
                            if type(custom_vars_payload[domain][selected_doc][cust_var])
                            == str
                            or (
                                "crm_bg"
                                in custom_vars_payload[domain][selected_doc][
                                    cust_var
                                ].keys()
                                and custom_vars_payload[domain][selected_doc][cust_var][
                                    "crm_bg"
                                ]
                                != "crm"
                                and custom_vars_payload[domain][selected_doc][cust_var][
                                    "crm_bg_option"
                                ]
                                != "array"
                            )
                        }
                        # custom_vars = list(custom_vars_dict.keys())
            # Doc Content
            is_file_exists = True
            doc_sections = {}
            if is_file_exists:
                json_db = {}
                master_db = master_mappings.find_one(
                    format_data(selected_doc, domain)["filter"]
                )
                if master_db != None:
                    json_db = dict(master_db)
                    json_db = bson_to_json(
                        "master_mappings", json_db, selected_doc, domain
                    )

                if json_db != None and domain in json_db.keys():
                    if selected_doc in list(json_db[domain].keys()):
                        doc_sections = json_db[domain][selected_doc]

            if len(list(doc_sections.keys())) > 0:
                for section_key in doc_sections.keys():
                    section_title = doc_sections[section_key]["section"].split(" ")
                    cur_key = section_key.split("_")
                    cur_section = {
                        "section": section_title[0],
                        "title": " ".join(section_title[1:]),
                        "parameter": cur_key[0],
                        "dependency": cur_key[1],
                        "value": cur_key[2],
                        "bgFunc": doc_sections[section_key]["domain_type"],
                        "document": doc_sections[section_key]["doc_name"],
                        "description": doc_sections[section_key]["doc_desc"],
                        "access": doc_sections[section_key]["access_level"],
                        "coOwner": doc_sections[section_key]["co_owner"],
                        "jsonConfig": doc_sections[section_key]["config"],
                    }
                    doc_content.append(cur_section)

            initial_key = list(doc_sections.keys())[0]
            initial_section = doc_sections[initial_key]["section"].split(" ")[0]
            initial_content = doc_sections[initial_key]["content"]
            uploaded_on = doc_sections[initial_key]["uploaded_on"]
            uploaded_by = doc_sections[initial_key]["uploaded_by"]

        return {
            "status": "Success!",
            "doc_content": doc_content,
            "ini_cont": initial_content,
            "ini_sec": initial_section,
            "custom_vars": custom_vars_dict,
            "uploaded_on": uploaded_on,
            "uploaded_by": uploaded_by,
        }
    except Exception as exp:
        return {"status": f"Error occured as {exp}"}


@application.route("/get_sel_sec_contents", methods=["GET", "POST"])
def get_sel_sec_contents():
    try:
        props = request.get_json()
        selected_doc = props["filename"]
        domain = props["domain"]
        user = props["user"]
        selected_sections = props["selected_sections"]
        doc_content = []
        final_sel_sections = []

        if len(selected_doc) > 0:
            is_file_exists = True
            if is_file_exists:
                sections_json_db = {}
                sections_db = sections_mappings.find_one(
                    format_data(selected_doc, None, user)["filter"]
                )
                if sections_db != None:
                    sections_json_db = dict(sections_db)
                    sections_json_db = bson_to_json(
                        "sections_mappings", sections_json_db, selected_doc, None, user
                    )

            doc_sections = {}
            if is_file_exists:
                json_db = {}
                master_db = master_mappings.find_one(
                    format_data(selected_doc, domain)["filter"]
                )
                if master_db != None:
                    json_db = dict(master_db)
                    json_db = bson_to_json(
                        "master_mappings", json_db, selected_doc, domain
                    )

                if json_db != None and domain in json_db.keys():
                    if selected_doc in list(json_db[domain].keys()):
                        doc_sections = json_db[domain][selected_doc]

            if len(list(doc_sections.keys())) > 0:
                # Duplicate Sections
                dup_sections = [
                    {
                        "dup_section": doc_sections[section_key]["section"],
                        "content": doc_sections[section_key]["content"],
                    }
                    for section_key in doc_sections.keys()
                    if re.search(
                        r"0.[a-z]$", doc_sections[section_key]["section"].split(" ")[0]
                    )
                ]
                if len(dup_sections) > 0:
                    for dup_sec in dup_sections:
                        cur_sec = dup_sec["dup_section"].split(" ")[0]
                        sec_ls = cur_sec.split(".")
                        parent_sec = ".".join(sec_ls[0 : len(sec_ls) - 2])
                        for section_key in doc_sections.keys():
                            section_title = doc_sections[section_key]["section"]
                            section_ = section_title.split(" ")[0]
                            if section_ == parent_sec:
                                sections_json_db[selected_doc][section_title][
                                    "content"
                                ] += dup_sec["content"]

                for section_key in doc_sections.keys():
                    section_title = doc_sections[section_key]["section"]
                    for sel_section in selected_sections:
                        # Add Parent Section if Duplicate is selected here!
                        # Add Updated Content
                        if section_title == sel_section and (
                            not re.search(r"0.[a-z]$", section_title.split(" ")[0])
                        ):
                            if (
                                sections_json_db != None
                                and selected_doc in sections_json_db.keys()
                                and section_title
                                in sections_json_db[selected_doc].keys()
                            ):
                                doc_content.append(
                                    sections_json_db[selected_doc][section_title][
                                        "content"
                                    ]
                                )
                            else:
                                doc_content.append(doc_sections[section_key]["content"])
                            final_sel_sections.append(sel_section)

        return {
            "status": "Received!",
            "sel_sec_content": doc_content,
            "sel_sections": final_sel_sections,
        }
    except Exception as exp:
        return {"status": f"Error occured as {exp}"}


@application.route("/get_uploaded_section_content", methods=["GET", "POST"])
def get_uploaded_section_content():
    try:
        prop = request.get_json()
        section_ = prop["section"]
        filename = prop["filename"]
        domain = prop["domain"]
        user = prop["user"]
        doc_content = ""
        is_file_exists = True
        if is_file_exists:
            section_json = {}
            sections_db = sections_mappings.find_one(
                format_data(filename, None, user)["filter"]
            )
            if sections_db != None:
                section_json = dict(sections_db)
                section_json = bson_to_json(
                    "sections_mappings", section_json, filename, None, user
                )

            if filename in section_json.keys():
                sections = section_json[filename]
                if section_ in sections.keys():
                    content = sections[section_]["content"]
                    return {
                        "status": "Received!",
                        "content": content,
                        "section": section_,
                    }

            uploaded_sections_db = {}
            master_db = master_mappings.find_one(
                format_data(filename, domain)["filter"]
            )
            if master_db != None:
                uploaded_sections_db = dict(master_db)
                uploaded_sections_db = bson_to_json(
                    "master_mappings", uploaded_sections_db, filename, domain
                )

            if (
                domain in uploaded_sections_db.keys()
                and filename in uploaded_sections_db[domain].keys()
            ):
                doc_sections = uploaded_sections_db[domain][filename]
                if len(list(doc_sections.keys())) > 0:
                    for section_key in doc_sections.keys():
                        section_title = doc_sections[section_key]["section"]
                        if section_title == section_:
                            doc_content = doc_sections[section_key]["content"]

        return {"status": "Received!", "content": doc_content}
    except Exception as exp:
        return {"status": f"Error occured as {exp}"}


@application.route("/set_uploaded_section_content", methods=["GET", "POST"])
def set_uploaded_section_content():
    try:
        prop = request.get_json()
        section_ = prop["section"]
        filename = prop["filename"]
        content = prop["content"]
        user = prop["user"]

        is_file_exists = True
        if is_file_exists:
            section_json = {}
            sections_db = sections_mappings.find_one(
                format_data(filename, None, user)["filter"]
            )
            if sections_db != None:
                section_json = dict(sections_db)
                section_json = bson_to_json(
                    "sections_mappings", section_json, filename, None, user
                )

            if filename in section_json.keys():
                sections = section_json[filename]
                if section_ in sections.keys():
                    sections[section_]["content"] = content
                else:
                    sections[section_] = {"content": content}
            else:
                section_json[filename] = {section_: {"content": content}}
            sections_bson = json_to_bson(
                {"sections": section_json[filename]},
                format_data(filename, None, user)["filter"],
            )
            if bool(sections_db):
                sections_mappings.update_one(
                    format_data(filename, None, user)["filter"],
                    {"$set": {"sections": section_json[filename]}},
                )
            else:
                sections_mappings.insert_many(sections_bson)

        return {"status": "Received!"}
    except Exception as exp:
        return {"status": f"Error occured as {exp}"}


@application.route("/get_word_file", methods=["GET", "POST"])
def create_word_doc():
    try:
        prop = request.get_json()
        conv_html = prop["conv_html"]
        conv_to = prop["conv_to"]
        domain = prop["domain"]
        filename = prop["filename"]
        is_crm = prop["is_crm"]
        crmid = prop["crmid"]

        # is_manual_gen = False

        # is_file_exists = True
        # if is_file_exists:
        #     json_db = {}
        #     master_db = master_mappings.find_one(
        #         format_data(filename, domain)["filter"]
        #     )
        #     if master_db != None:
        #         json_db = dict(master_db)
        #         json_db = bson_to_json("master_mappings", json_db, filename, domain)

        # if json_db != None:
        #     if domain in json_db.keys() and filename in json_db[domain].keys():
        #         is_manual_gen = (
        #             json_db[domain][filename][
        #                 list(json_db[domain][filename].keys())[0]
        #             ]["config"]
        #             != "bg_config"
        #         )
        # Custom vars mappings
        custom_vars_dict = {}
        is_file_exists = True
        if is_file_exists:
            custom_vars_db = custom_vars_mappings.find_one(
                format_data(filename, domain)["filter"]
            )
            if custom_vars_db != None:
                json_db = dict(custom_vars_db)
                json_db = bson_to_json(
                    "custom_vars_mappings", json_db, filename, domain
                )

            if json_db != None:
                if domain in json_db.keys() and filename in json_db[domain].keys():
                    custom_vars_dict = json_db[domain][filename]

        if custom_vars_dict != None and len(custom_vars_dict.keys()) > 0:
            custom_vars = list(custom_vars_dict.keys())
        else:
            custom_vars = []
        # Get BBT Data
        if is_crm:
            bbt_data = get_bbt_data(crmid)
            if len(bbt_data.keys()) > 0:
                pass
            else:
                raise Exception("Invalid CRM-ID!")
        else:
            bbt_data = {
                "customer": "Test customer",
                "market": "Test market",
                "opportunityname": "Test opportunity name",
                "country": "Test country",
                "region": "Test region",
                "subregion": "Test subregion",
            }

        bbt_data["crmid"] = crmid

        # Update images in custom-vars
        for cvar in custom_vars:
            cvar_pat = re.sub("\[", "\[", cvar)
            cvar_pat = re.sub("\]", "\]", cvar_pat)
            if len(re.findall(cvar_pat, conv_html)) > 0:
                if (
                    type(custom_vars_dict[cvar]) != str
                    and custom_vars_dict[cvar]["crm_bg_option"] == "img"
                    and cvar in prop.keys()
                    and isBase64(prop[cvar].split("base64,")[1])
                ):
                    conv_html = conv_html.replace(
                        cvar,
                        f"<img src='{prop[cvar]}'/>",
                    )
        parser = DocHTMLParser()
        parser.set_document()
        parser.set_fig_count()
        parser.set_table_count()
        parser.set_prev_start_tag()
        parser.set_table_tags()
        parser.feed(conv_html)
        document = parser.get_document()
        for para in document.paragraphs:
            if len(re.findall("Figure [0-9]*:", para.text)) > 0:
                fig_tab_caption(para, True)
            elif len(re.findall("Table [0-9]*:", para.text)) > 0:
                fig_tab_caption(para)
        document.save("doc_assist_be/services/converted_doc.docx")
        document = Document("doc_assist_be/services/converted_doc.docx")
        for para in document.paragraphs:
            if para.text == "%TOC%":
                para.text = ""
                run = para.add_run()
                add_list_of_table(run)
            elif para.text == "%TOF%":
                para.text = ""
                run = para.add_run()
                add_list_of_fig_tab(run, True)
            elif para.text == "%TOT%":
                para.text = ""
                run = para.add_run()
                add_list_of_fig_tab(run)
            else:
                update_custom_vars(
                    bbt_data, custom_vars_dict, custom_vars, para, prop, -1, -1
                )

        # table = document.tables[0]
        for table in document.tables:
            for row_idx, each_row in enumerate(table.rows):
                for each_cell in each_row.cells:
                    start_idx = 0
                    if (
                        len(each_cell.paragraphs[0].text) == 0
                        and len(each_cell.paragraphs) > 1
                    ):
                        delete_paragraph(each_cell.paragraphs[0])
                        # start_idx = 1
                    if row_idx == 0:
                        if len(each_cell.paragraphs) > 0:
                            for para in each_cell.paragraphs[start_idx:]:
                                update_custom_vars(
                                    bbt_data,
                                    custom_vars_dict,
                                    custom_vars,
                                    para,
                                    prop,
                                    -1,
                                    -1,
                                )
                                format_table(each_cell, para, True)
                        set_table_header_bg_color(each_cell, "#005AFF")
                    else:
                        if len(each_cell.paragraphs) > 0:
                            for para in each_cell.paragraphs[start_idx:]:
                                update_custom_vars(
                                    bbt_data,
                                    custom_vars_dict,
                                    custom_vars,
                                    para,
                                    prop,
                                    -1,
                                    -1,
                                )

                        #         format_table(each_cell, para, False)
                        set_table_header_bg_color(each_cell, "#CCCCCC")
                    set_cell_border(
                        each_cell,
                        top={"sz": 12, "color": "#FFFFFF", "val": "single"},
                        bottom={"sz": 12, "color": "#FFFFFF", "val": "single"},
                        start={"sz": 12, "color": "#FFFFFF", "val": "single"},
                        end={"sz": 12, "color": "#FFFFFF", "val": "single"},
                    )
        document.save("doc_assist_be/services/converted_doc.docx")
        # pythoncom.CoInitialize()
        # word = comtypes.client.CreateObject(
        #     "Word.Application"
        # )  # opens the word application
        # doc = word.Documents.Open(filepath)  # opens the specified file
        # res = doc.Fields.Update()  # updates field, returns 0 if successful
        # print(res)
        # doc.Save()
        # doc.Close()
        # word.Quit()
        if conv_to == "pdf":
            try:
                convert(filepath)
            finally:
                return send_file("doc_assist_be/services/converted_doc.pdf")
        return send_file("doc_assist_be/services/converted_doc.docx")
    except Exception as exp:
        return {"status": f"Error occured as {exp}"}


@application.route("/matched_sections", methods=["GET", "POST"])
def matched_sections():
    if len(request.get_json()) > 0:
        prop = request.get_json()
        filename = prop["filename"]
        sections = prop["match_sections"]
        content = prop["content"]
        domain = prop["domain"]
        user = prop["user"]
        # Extract Sections
        sections = list(
            map(
                lambda sec: {
                    "id": sec["id"],
                    "value": " ".join(
                        [
                            re.sub(r"[.]$", "", sec["value"].split(" ")[0]),
                            " ".join(sec["value"].split(" ")[1:]),
                        ]
                    ),
                },
                sections,
            ),
        )
        sections = list(
            filter(
                lambda sec: not sec["value"]
                .split(" ")[0]
                .lower()
                .__contains__("figure")
                and not sec["value"].split(" ")[0].lower().__contains__("table"),
                sections,
            )
        )
        next_avail_ver = list(
            filter(
                lambda sec: sec["value"].split(" ")[0].isnumeric(),
                sections,
            )
        )
        next_avail_ver = int(
            next_avail_ver[len(next_avail_ver) - 1]["value"].split(" ")[0]
        )
        for section in sections:
            value = section["value"]
            version = value.split(" ")[0]
            if version.lower().__contains__("figure") or version.lower().__contains__(
                "table"
            ):
                version = list(filter(len, value.split(":")))[0]
                section["section"] = version.strip()
                section["value"] = " ".join(value.split(":")[1:]).strip()
            elif version.split(".")[0].isnumeric():
                re_version = ".".join(list(filter(len, version.split("."))))
                section["section"] = re_version.strip()
                section["value"] = " ".join(value.split(" ")[1:])
            else:
                re_version = str(next_avail_ver + 1)
                section["section"] = re_version.strip()
                section["value"] = value

        # with open(sections_mappings, "r") as outfile:
        #     section_json = json.load(outfile)

        section_json_db = sections_mappings.find_one(
            format_data(filename, None, user)["filter"]
        )
        if section_json_db == None:
            section_json = {}
        else:
            section_json = dict(section_json_db)
            section_json = bson_to_json(
                "sections_mappings", section_json, filename, None, user
            )
        # if(filename not in section_json.keys()):
        section_json[filename] = {}

        # Create hierarchy based on version with content, section and children
        version_hierarchy = {}

        # Custom-variables mapping
        custom_vars = {}

        for section in sections:
            version = section["section"]
            node = {"children": []}
            sub_vers = version.split(".")
            node_ver = ".".join(sub_vers)
            section_json[filename][node_ver] = {
                "parameter": "",
                "value": "",
                "dependency": "",
                "title": section["value"],
                "id": section["id"],
            }
            if section["id"] in content.keys():
                # Custom variables extraction
                match_ls = re.findall(custom_vars_pat, content[section["id"]])
                if len(match_ls) > 0:
                    for cust_var in match_ls:
                        if cust_var not in custom_vars.keys():
                            custom_vars[cust_var] = ""
                section_json[filename][node_ver]["content"] = content[section["id"]]
            else:
                section_json[filename][node_ver]["content"] = ""
            if len(sub_vers) > 1:
                parent = ".".join(sub_vers[0 : len(sub_vers) - 1])
                if parent not in version_hierarchy.keys():
                    version_hierarchy[parent] = {"children": []}
                version_hierarchy[parent]["children"].append(node_ver)
                version_hierarchy[node_ver] = node
                # print(f"Parent: {parent} - Child: {version}")
            else:
                version_hierarchy[node_ver] = node
        # Restricted List
        restricted_keywrds_ls = get_restricted_keywrds_list(restricted_ls)
        restricted_keywrds_sections = get_sensitive_sections(
            restricted_keywrds_ls, content, sections
        )
        # with open(sections_mappings, "w") as outfile:
        #     json.dump(section_json, outfile)
        sections_bson = json_to_bson(
            {"sections": section_json[filename]},
            format_data(filename, None, user)["filter"],
        )
        if bool(section_json_db):
            sections_mappings.update_one(
                format_data(filename, None, user)["filter"],
                {"$set": {"sections": section_json[filename]}},
            )
        else:
            sections_mappings.insert_many(sections_bson)

        # Save Custom Vars
        is_file_exists = True
        if is_file_exists:
            # custom_vars_payload = {domain: {filename: custom_vars}}
            custom_vars_bson = json_to_bson(
                {"custom_vars": custom_vars}, format_data(filename, domain)["filter"]
            )
            custom_vars_mappings.insert_many(custom_vars_bson)
            # with open(custom_vars_mappings, "w") as outfile:
            #     json.dump(custom_vars_payload, outfile)

        is_file_exists = True
        # payload = {filename: version_hierarchy}
        if is_file_exists:
            heirarchy_bson = json_to_bson(
                {"version_heirarchy": version_hierarchy},
                format_data(filename, None, user)["filter"],
            )
            heirarchy_mappings.insert_many(heirarchy_bson)
            # with open(heirarchy_mappings, "w") as outfile:
            #     json.dump(payload, outfile)
        ini_content = section_json[filename][list(section_json[filename].keys())[0]][
            "content"
        ]
        ini_section = list(version_hierarchy.keys())[0]
        return {
            "status": "Received!",
            "extracted_sections": sections,
            "ini_content": ini_content,
            "ini_section": ini_section,
            "custom_vars": list(custom_vars.keys()),
            "sensitive_keywords_sections": restricted_keywrds_sections,
            "restricted_keywrds_pat": restricted_keywrds_ls,
        }
    # except Exception as exp:
    #     return {"status": f"Error occured as {exp}"}


@application.route("/get_section_content", methods=["GET", "POST"])
def get_section_content():
    try:
        prop = request.get_json()
        section_ = prop["section"]
        filename = prop["filename"]
        user = prop["user"]
        is_file_exists = True
        # Section DB
        if is_file_exists:
            section_json = {}
            sections_db = sections_mappings.find_one(
                format_data(filename, None, user)["filter"]
            )
            if sections_db != None:
                section_json = dict(sections_db)
                section_json = bson_to_json(
                    "sections_mappings", section_json, filename, None, user
                )

            if filename in section_json.keys():
                sections = section_json[filename]
                if section_ in sections.keys():
                    content = sections[section_]["content"]
                    parameter = sections[section_]["parameter"]
                    dependency = sections[section_]["dependency"]
                    value = sections[section_]["value"]
                    title = sections[section_]["title"]
                    return {
                        "status": "Received!",
                        "content": content,
                        "parameter": parameter,
                        "dependency": dependency,
                        "value": value,
                        "section": section_,
                        "title": title,
                    }

        return {"status": "Received!", "content": ""}
    except Exception as exp:
        return {"status": f"Error occured as {exp}"}


@application.route("/set_section_content", methods=["GET", "POST"])
def set_section_content():
    try:
        prop = request.get_json()
        domain = prop["domain"]
        user = prop["user"]
        section_ = prop["section"]
        filename = prop["filename"]
        content = prop["content"]
        parameter = prop["parameter"]
        dependency = prop["dependency"]
        value = prop["value"]
        title = prop["title"]

        # Custom Vars
        # is_file_exists = os.path.isfile(custom_vars_mappings)
        is_file_exists = True
        custom_vars_payload = {}
        custom_vars = {}
        if is_file_exists:
            custom_vars_db = custom_vars_mappings.find_one(
                format_data(filename, domain)["filter"]
            )
            if custom_vars_db != None:
                custom_vars_payload = dict(custom_vars_db)
                custom_vars_payload = bson_to_json(
                    "custom_vars_mappings", custom_vars_payload, filename, domain
                )
            # with open(custom_vars_mappings, "r") as outfile:
            #     custom_vars_payload = json.load(outfile)

        if (
            domain in custom_vars_payload.keys()
            and filename in custom_vars_payload[domain].keys()
        ):
            custom_vars = custom_vars_payload[domain][filename]

        # is_file_exists = os.path.isfile(sections_mappings)
        if is_file_exists:
            section_json = {}
            sections_db = sections_mappings.find_one(
                format_data(filename, None, user)["filter"]
            )
            if sections_db != None:
                section_json = dict(sections_db)
                section_json = bson_to_json(
                    "sections_mappings", section_json, filename, None, user
                )
            # with open(sections_mappings, "r") as outfile:
            #     section_json = json.load(outfile)

            if filename in section_json.keys():
                sections = section_json[filename]
                if section_ not in sections.keys():
                    sections[section_] = {}

                match_ls = re.findall(custom_vars_pat, content)
                if len(match_ls) > 0:
                    for cust_var in match_ls:
                        if cust_var not in custom_vars.keys():
                            custom_vars[cust_var] = ""
                sections[section_]["content"] = content
                sections[section_]["parameter"] = parameter
                sections[section_]["dependency"] = dependency
                sections[section_]["value"] = value
                sections[section_]["title"] = title
            else:
                section_json[filename] = {
                    section_: {
                        "content": content,
                        "parameter": parameter,
                        "dependency": dependency,
                        "value": value,
                        "title": title,
                    }
                }
            # Custom Vars DB
            custom_vars_bson = json_to_bson(
                {"custom_vars": custom_vars}, format_data(filename, domain)["filter"]
            )
            if bool(custom_vars_db):
                custom_vars_mappings.update_one(
                    format_data(filename, domain)["filter"],
                    {"$set": {"custom_vars": custom_vars_payload[domain][filename]}},
                )
            else:
                custom_vars_mappings.insert_many(custom_vars_bson)

            # Section DB
            sections_bson = json_to_bson(
                {"sections": section_json[filename]},
                format_data(filename, None, user)["filter"],
            )
            if bool(sections_db):
                sections_mappings.update_one(
                    format_data(filename, None, user)["filter"],
                    {"$set": {"sections": section_json[filename]}},
                )
            else:
                sections_mappings.insert_many(sections_bson)
            # with open(custom_vars_mappings, "w") as outfile:
            #     json.dump(custom_vars_payload, outfile)

            # with open(sections_mappings, "w") as outfile:
            #     json.dump(section_json, outfile)

        return {"status": "Received!", "custom_vars": list(custom_vars.keys())}
    except Exception as exp:
        return {"status": f"Error occured as {exp}"}


@application.route("/all_sections_content", methods=["GET", "POST"])
def all_sections_content():
    try:
        prop = request.get_json()
        filename = prop["filename"]
        user = prop["user"]
        is_file_exists = True
        if is_file_exists:
            section_json = {}
            sections_db = sections_mappings.find_one(
                format_data(filename, None, user)["filter"]
            )
            if sections_db != None:
                section_json = dict(sections_db)
                section_json = bson_to_json(
                    "sections_mappings", section_json, filename, None, user
                )

            if filename in section_json.keys():
                sections = section_json[filename]
                all_sections_data = []
                for section_ in list(sections.keys()):
                    content = sections[section_]["content"]
                    parameter = sections[section_]["parameter"]
                    dependency = sections[section_]["dependency"]
                    value = sections[section_]["value"]
                    title = sections[section_]["title"]
                    all_sections_data.append(
                        {
                            "content": content,
                            "parameter": parameter,
                            "dependency": dependency,
                            "value": value,
                            "section": section_,
                            "title": title,
                        }
                    )

                return {"status": "Received!", "sections_data": all_sections_data}

        return {"status": "Received!", "sections_data": []}
    except Exception as exp:
        return {"status": f"Error occured as {exp}"}


@application.route("/add_new_sec", methods=["GET", "POST"])
def add_new_sec():
    try:
        prop = request.get_json()
        is_section = prop["is_section"]
        sec_header = prop["section_header"]
        filename = prop["filename"]
        cur_section = prop["cur_section"]
        user = prop["user"]
        is_sections_exists = True
        is_versions_exists = True

        avail_version = -1
        new_keys = {}
        if is_versions_exists:
            version_db = {}
            version_db_ = {filename: {}}
            version_h_db = heirarchy_mappings.find_one(
                format_data(filename, None, user)["filter"]
            )
            if version_h_db != None:
                version_db = dict(version_h_db)
                version_db = bson_to_json(
                    "heirarchy_mappings", version_db, filename, None, user
                )

            if filename in version_db.keys():
                sections = list(version_db[filename].keys())
                next_sec = cur_section.split(".")

                filtered_arr = []
                unfiltered_arr = []

                if is_section:
                    for sec in sections:
                        sec = sec.split(".")
                        if (
                            len(sec) >= len(next_sec)
                            and sec[0 : len(next_sec) - 1]
                            == next_sec[0 : len(next_sec) - 1]
                        ):
                            filtered_arr.append(".".join(sec))
                        else:
                            unfiltered_arr.append(".".join(sec))

                    last_digit_arr = []
                    for sec in filtered_arr:
                        sec = sec.split(".")
                        last_digit_arr.append(sec[len(next_sec) - 1])
                    for idx, ver in enumerate(last_digit_arr):
                        if int(ver) >= int(next_sec[len(next_sec) - 1]) + 1:
                            new_ver = filtered_arr[idx].split(".")
                            new_ver[len(next_sec) - 1] = str(int(ver) + 1)
                            new_keys[filtered_arr[idx]] = ".".join(new_ver)
                        else:
                            new_keys[filtered_arr[idx]] = filtered_arr[idx]
                    sections_ = []
                    for sec in sections:
                        if sec in new_keys.keys():
                            sections_.append(new_keys[sec])
                        else:
                            sections_.append(sec)
                    # Available Next Version
                    avail_version = next_sec
                    avail_version[len(next_sec) - 1] = str(
                        int(avail_version[len(next_sec) - 1]) + 1
                    )
                    avail_version = ".".join(avail_version)
                    sections_.append(avail_version)
                    # Create Latest Heirarchy
                    version_hierarchy = create_ver_heirarchy(sections_)
                    version_db_[filename] = version_hierarchy

                else:
                    version_hierarchy = create_ver_heirarchy(sections)
                    if cur_section in sections:
                        cur_section_children = version_db[filename][cur_section][
                            "children"
                        ]
                        # Sort versions if needed using natsort
                        if len(cur_section_children) > 0:
                            last_child = cur_section_children[
                                len(cur_section_children) - 1
                            ]
                            last_child_ls = last_child.split(".")
                            next_avail_child = (
                                int(last_child_ls[len(last_child_ls) - 1]) + 1
                            )
                            avail_version = f"{'.'.join(last_child_ls[0: len(last_child_ls)-1])}.{str(next_avail_child)}"
                        else:
                            avail_version = f"{cur_section}.1"

                        version_hierarchy[cur_section]["children"].append(avail_version)
                        version_hierarchy[avail_version] = {"children": []}
                        version_db_[filename] = version_hierarchy

                # Versions DB
                versions_bson = json_to_bson(
                    {"version_heirarchy": version_db_[filename]},
                    format_data(filename, None, user)["filter"],
                )
                if bool(version_h_db):
                    heirarchy_mappings.update_one(
                        format_data(filename, None, user)["filter"],
                        {"$set": {"version_heirarchy": version_db_[filename]}},
                    )
                else:
                    heirarchy_mappings.insert_many(versions_bson)

        if is_sections_exists:
            sections = {}
            sections_ = {filename: {}}
            sections_db = sections_mappings.find_one(
                format_data(filename, None, user)["filter"]
            )
            if sections_db != None:
                sections = dict(sections_db)
                sections = bson_to_json(
                    "sections_mappings", sections, filename, None, user
                )

            if filename in sections.keys():
                cur_sections = sections[filename]
                for sec in cur_sections.keys():
                    if sec in new_keys.keys():
                        sections_[filename][new_keys[sec]] = cur_sections[sec]
                    else:
                        sections_[filename][sec] = cur_sections[sec]
                sections_[filename][avail_version] = {}
                sections_[filename][avail_version]["content"] = ""
                sections_[filename][avail_version]["parameter"] = ""
                sections_[filename][avail_version]["dependency"] = ""
                sections_[filename][avail_version]["value"] = ""
                sections_[filename][avail_version]["title"] = sec_header
                sections_[filename][avail_version]["id"] = f"#_TOC_{sec_header}"
                # Section DB
                sections_bson = json_to_bson(
                    {"sections": sections_[filename]},
                    format_data(filename, None, user)["filter"],
                )
                if bool(sections_db):
                    sections_mappings.update_one(
                        format_data(filename, None, user)["filter"],
                        {"$set": {"sections": sections_[filename]}},
                    )
                else:
                    sections_mappings.insert_many(sections_bson)
                # with open(sections_mappings, "w") as outfile:
                #     json.dump(sections_, outfile)

                new_sections = []
                cur_sections = sections_[filename]
                for sec_header, sec_value in cur_sections.items():
                    new_sections.append(
                        {
                            "section": sec_header,
                            "value": sec_value["title"],
                            "id": sec_value["id"],
                        }
                    )

                new_sections = natsorted(
                    new_sections, key=lambda section: section["section"]
                )
                ini_content = cur_sections[list(cur_sections.keys())[0]]["content"]
                ini_section = list(version_db[filename].keys())[0]

            return {
                "status": "Received!",
                "avail_section": avail_version,
                "sections": new_sections,
                "ini_content": ini_content,
                "ini_section": ini_section,
            }
    except Exception as exp:
        return {"status": f"Error occured as {exp}"}


@application.route("/del_sec", methods=["GET", "POST"])
def del_sec():
    try:
        prop = request.get_json()
        filename = prop["filename"]
        cur_section = prop["cur_section"]
        user = prop["user"]
        is_sections_exists = True
        is_versions_exists = True

        if is_versions_exists and is_sections_exists:
            versions_db = {}
            version_h_db = heirarchy_mappings.find_one(
                format_data(filename, None, user)["filter"]
            )
            if version_h_db != None:
                versions_db = dict(version_h_db)
                versions_db = bson_to_json(
                    "heirarchy_mappings", versions_db, filename, None, user
                )

            sections_db = {}
            section_db = sections_mappings.find_one(
                format_data(filename, None, user)["filter"]
            )
            if section_db != None:
                sections_db = dict(section_db)
                sections_db = bson_to_json(
                    "sections_mappings", sections_db, filename, None, user
                )

            if filename in versions_db and cur_section in versions_db[filename]:
                # print(list(versions_db[filename].keys()))
                all_sections = list(
                    filter(
                        lambda section: ".".join(
                            section.split(".")[0 : len(cur_section.split("."))]
                        )
                        == cur_section,
                        list(versions_db[filename].keys()),
                    )
                )
                # print(all_sections)
                for section in all_sections:
                    del versions_db[filename][section]
                    del sections_db[filename][section]
            # Version DB
            versions_bson = json_to_bson(
                {"version_heirarchy": versions_db[filename]},
                format_data(filename, None, user)["filter"],
            )
            if bool(version_h_db):
                heirarchy_mappings.update_one(
                    format_data(filename, None, user)["filter"],
                    {"$set": {"version_heirarchy": versions_db[filename]}},
                )
            else:
                heirarchy_mappings.insert_many(versions_bson)

            # Sections DB
            sections_bson = json_to_bson(
                {"sections": sections_db[filename]},
                format_data(filename, None, user)["filter"],
            )
            if bool(sections_db):
                sections_mappings.update_one(
                    format_data(filename, None, user)["filter"],
                    {"$set": {"sections": sections_db[filename]}},
                )
            else:
                sections_mappings.insert_many(sections_bson)

            new_sections = []
            cur_sections = sections_db[filename]
            for sec_header, sec_value in cur_sections.items():
                new_sections.append(
                    {
                        "section": sec_header,
                        "value": sec_value["title"],
                        "id": sec_value["id"],
                    }
                )

            new_sections = natsorted(
                new_sections, key=lambda section: section["section"]
            )
            ini_content = cur_sections[list(cur_sections.keys())[0]]["content"]
            ini_section = list(versions_db[filename].keys())[0]

            return {
                "status": "Received!",
                "sections": new_sections,
                "ini_content": ini_content,
                "ini_section": ini_section,
            }
    except Exception as exp:
        return {"status": f"Error occured as {exp}"}


@application.route("/dup_sec", methods=["GET", "POST"])
def dup_sec():
    try:
        prop = request.get_json()
        filename = prop["filename"]
        cur_section = prop["cur_section"]
        user = prop["user"]
        is_sections_exists = True
        is_versions_exists = True

        if is_sections_exists:
            sections = {}
            sections_db = sections_mappings.find_one(
                format_data(filename, None, user)["filter"]
            )
            if sections_db != None:
                sections = dict(sections_db)
                sections = bson_to_json(
                    "sections_mappings", sections, filename, None, user
                )

            if filename in sections.keys():
                cur_sections = sections[filename]
                # print(cur_sections.keys(), cur_section)
                if cur_section in cur_sections.keys():
                    pass
                else:
                    return "Please add parent section to create duplicate!"

        avail_version = -1
        if is_versions_exists:
            version_db = {}
            version_h_db = heirarchy_mappings.find_one(
                format_data(filename, None, user)["filter"]
            )
            if version_h_db != None:
                version_db = dict(version_h_db)
                version_db = bson_to_json(
                    "heirarchy_mappings", version_db, filename, None, user
                )
            if filename in version_db.keys():
                sections = list(version_db[filename].keys())

                if cur_section in sections:
                    cur_section_children = version_db[filename][cur_section]["children"]
                    # Sort versions if needed using natsort
                    if len(cur_section_children) > 0:
                        cur_sec_dup_children = [
                            child
                            for child in cur_section_children
                            if re.search(r"0.[a-z]$", child)
                        ]
                        if len(cur_sec_dup_children) > 0:
                            last_child = max(cur_sec_dup_children)
                            last_child_ls = last_child.split(".")

                            next_avail_child = (
                                ord(last_child_ls[len(last_child_ls) - 1]) + 1
                            )
                            avail_version = f"{'.'.join(last_child_ls[0: len(last_child_ls)-1])}.{chr(next_avail_child)}"
                        else:
                            avail_version = f"{cur_section}.0.a"
                    else:
                        avail_version = f"{cur_section}.0.a"

                    cur_section_children.append(avail_version)

                avail_version = str(avail_version)
                version_db[filename][avail_version] = {"children": []}

                versions_bson = json_to_bson(
                    {"version_heirarchy": version_db[filename]},
                    format_data(filename, None, user)["filter"],
                )
                if bool(version_h_db):
                    heirarchy_mappings.update_one(
                        format_data(filename, None, user)["filter"],
                        {"$set": {"version_heirarchy": version_db[filename]}},
                    )
                else:
                    heirarchy_mappings.insert_many(versions_bson)

        if is_sections_exists:
            sections = {}
            sections_db = sections_mappings.find_one(
                format_data(filename, None, user)["filter"]
            )
            if sections_db != None:
                sections = dict(sections_db)
                sections = bson_to_json(
                    "sections_mappings", sections, filename, None, user
                )

            if filename in sections.keys():
                cur_sections = sections[filename]
                cur_sections[avail_version] = {}
                cur_sections[avail_version]["content"] = cur_sections[cur_section][
                    "content"
                ]
                cur_sections[avail_version]["parameter"] = ""
                cur_sections[avail_version]["dependency"] = ""
                cur_sections[avail_version]["value"] = ""
                cur_sections[avail_version][
                    "title"
                ] = f"Duplicate of {cur_sections[cur_section]['title']}"
                cur_sections[avail_version]["id"] = f"#_TOC_{avail_version}"

                sections_bson = json_to_bson(
                    {"sections": sections[filename]},
                    format_data(filename, None, user)["filter"],
                )
                if bool(sections_db):
                    sections_mappings.update_one(
                        format_data(filename, None, user)["filter"],
                        {"$set": {"sections": sections[filename]}},
                    )
                else:
                    sections_mappings.insert_many(sections_bson)

                new_sections = []
                for sec_header, sec_value in cur_sections.items():
                    new_sections.append(
                        {
                            "section": sec_header,
                            "value": sec_value["title"],
                            "id": sec_value["id"],
                        }
                    )

                new_sections = natsorted(
                    new_sections, key=lambda section: section["section"]
                )
                ini_content = cur_sections[list(cur_sections.keys())[0]]["content"]
                ini_section = list(version_db[filename].keys())[0]

            return {
                "status": "Received!",
                "avail_section": avail_version,
                "sections": new_sections,
                "ini_content": ini_content,
                "ini_section": ini_section,
            }
    except Exception as exp:
        return {"status": f"Error occured as {exp}"}


# Download history
@application.route("/set_download_history", methods=["GET", "POST"])
def set_download_history():
    try:
        prop = request.get_json()
        master_db_filename = prop["filename"]
        filename = prop["newFilename"]
        user = prop["user"]
        domain = prop["domain"]
        sections = prop["sections"]
        sec_contents = prop["contents"]
        all_sections = prop["all_sections"]
        doc_name = ""
        is_file_exists = True
        if is_file_exists:
            json_db = {}
            master_db = master_mappings.find_one(
                format_data(master_db_filename, domain)["filter"]
            )
            if master_db != None:
                json_db = dict(master_db)
                json_db = bson_to_json(
                    "master_mappings", json_db, master_db_filename, domain
                )
                if (
                    json_db != None
                    and domain in json_db.keys()
                    and master_db_filename in json_db[domain].keys()
                ):
                    section_name = list(json_db[domain][master_db_filename].keys())[0]
                    doc_name = json_db[domain][master_db_filename][section_name][
                        "doc_name"
                    ]
                    uploaded_by = json_db[domain][master_db_filename][section_name][
                        "uploaded_by"
                    ]

        sections = [section.split(" ")[0] for section in sections]
        download_history = {}
        for sec_idx, section in enumerate(all_sections):
            if section["section"] in sections:
                download_history[
                    f"{section['parameter']}_{section['dependency']}_{section['paramValue']}"
                ] = {
                    "content": sec_contents[sec_idx],
                    "section": f"{section['section']} {section['value']}",
                    "doc_name": doc_name,
                    "uploaded_on": str(datetime.now()),
                    "uploaded_by": uploaded_by,
                }
        final_json = {user: {domain: {filename: download_history}}}

        if is_file_exists:
            json_db = {}
            history_db = download_history_mappings.find_one(
                format_data(filename, domain, user)["filter"]
            )
            if history_db != None:
                json_db = dict(history_db)
                json_db = bson_to_json(
                    "download_history_mappings", json_db, filename, domain, user
                )

            if json_db != None:
                # new_param_obj = {}
                if user in json_db.keys() and domain in json_db[user].keys():
                    json_db[user][domain][filename] = final_json[user][domain][filename]
                elif user in json_db.keys():
                    json_db[user].update(final_json[user])
                    history_bson = json_to_bson(
                        {"mappings": json_db[user][domain][filename]},
                        format_data(filename, domain, user)["filter"],
                    )
                    if bool(history_db):
                        download_history_mappings.update_one(
                            format_data(filename, domain, user)["filter"],
                            {"$set": {"mappings": json_db[user][domain][filename]}},
                        )
                    else:
                        download_history_mappings.insert_many(history_bson)
                else:
                    json_db[user] = final_json[user]
                history_bson = json_to_bson(
                    {"mappings": json_db[user][domain][filename]},
                    format_data(filename, domain, user)["filter"],
                )
                if bool(history_db):
                    download_history_mappings.update_one(
                        format_data(filename, domain, user)["filter"],
                        {"$set": {"mappings": json_db[user][domain][filename]}},
                    )
                else:
                    download_history_mappings.insert_many(history_bson)
            else:
                history_bson = json_to_bson(
                    {"mappings": final_json[user][domain][filename]},
                    format_data(filename, domain, user)["filter"],
                )
                if bool(history_db):
                    download_history_mappings.update_one(
                        format_data(filename, domain, user)["filter"],
                        {"$set": {"mappings": final_json[user][domain][filename]}},
                    )
                else:
                    download_history_mappings.insert_many(history_bson)
        return {"status": "Received!"}
    except Exception as exp:
        return {"status": f"Error occured as {exp}"}


@application.route("/get_download_history", methods=["GET", "POST"])
def get_download_history():
    try:
        prop = request.get_json()
        domain = prop["domain"]
        user = prop["user"]
        props = []
        is_file_exists = True
        if is_file_exists:
            json_db = {}
            history_db = download_history_mappings.find(
                {"user_domain": domain, "user_name": user}
            )
            if history_db != None:
                history_db = list(history_db)
                for file in history_db:
                    json_db[file["user_file"]] = file["mappings"]
                json_db = {user: {domain: json_db}}

            if json_db != None:
                if user in json_db.keys():
                    if domain in json_db[user].keys():
                        for doc in json_db[user][domain].keys():
                            doc_param = list(json_db[user][domain][doc].keys())[0]
                            doc_name = json_db[user][domain][doc][doc_param]["doc_name"]
                            doc_desc = json_db[user][domain][doc][doc_param][
                                "uploaded_on"
                            ]
                            file_uploaded = doc
                            props.append(
                                {
                                    "doc_name": doc_name,
                                    "doc_desc": doc_desc,
                                    "file_uploaded": file_uploaded,
                                }
                            )

        return {"status": "Success!", "docs": props}
    except Exception as exp:
        return {"status": f"Error occured as {exp}"}


@application.route("/get_dh_sections", methods=["GET", "POST"])
def get_dh_sections():
    try:
        prop = request.get_json()
        selected_doc = prop["selectedDoc"]
        domain = prop["domain"]
        user = prop["user"]
        doc_content = []
        initial_content = ""
        if len(selected_doc) > 0:
            is_file_exists = True
            doc_sections = {}
            if is_file_exists:
                json_db = {}
                history_db = download_history_mappings.find_one(
                    format_data(selected_doc, domain, user)["filter"]
                )
                if history_db != None:
                    json_db = dict(history_db)
                    json_db = bson_to_json(
                        "download_history_mappings", json_db, selected_doc, domain, user
                    )

                if (
                    json_db != None
                    and user in json_db.keys()
                    and domain in json_db[user].keys()
                ):
                    if selected_doc in list(json_db[user][domain].keys()):
                        doc_sections = json_db[user][domain][selected_doc]

            if len(list(doc_sections.keys())) > 0:
                for section_key in doc_sections.keys():
                    section_title = doc_sections[section_key]["section"].split(" ")
                    cur_key = section_key.split("_")
                    cur_section = {
                        "section": section_title[0],
                        "title": " ".join(section_title[1:]),
                        "parameter": cur_key[0],
                        "dependency": cur_key[1],
                        "value": cur_key[2],
                    }
                    doc_content.append(cur_section)

            initial_key = list(doc_sections.keys())[0]
            initial_section = doc_sections[initial_key]["section"].split(" ")[0]
            initial_content = doc_sections[initial_key]["content"]
            uploaded_on = doc_sections[initial_key]["uploaded_on"]
            uploaded_by = doc_sections[initial_key]["uploaded_by"]

        return {
            "status": "Success!",
            "doc_content": doc_content,
            "ini_cont": initial_content,
            "ini_sec": initial_section,
            "uploaded_on": uploaded_on,
            "uploaded_by": uploaded_by,
        }
    except Exception as exp:
        return {"status": f"Error occured as {exp}"}


@application.route("/get_sel_sec_dh_contents", methods=["GET", "POST"])
def get_sel_sec_dh_contents():
    try:
        props = request.get_json()
        selected_doc = props["filename"]
        domain = props["domain"]
        selected_sections = props["selected_sections"]
        user = props["user"]
        doc_content = []

        if len(selected_doc) > 0:
            is_file_exists = True
            if is_file_exists:
                sections_json_db = {}
                sections_db = sections_mappings.find_one(
                    format_data(selected_doc, None, user)["filter"]
                )
                if sections_db != None:
                    sections_json_db = dict(sections_db)
                    sections_json_db = bson_to_json(
                        "sections_mappings", sections_json_db, selected_doc, None, user
                    )

            doc_sections = {}
            if is_file_exists:
                json_db = {}
                history_db = download_history_mappings.find_one(
                    format_data(selected_doc, domain, user)["filter"]
                )
                if history_db != None:
                    json_db = dict(history_db)
                    json_db = bson_to_json(
                        "download_history_mappings", json_db, selected_doc, domain, user
                    )

                if (
                    json_db != None
                    and user in json_db.keys()
                    and domain in json_db[user].keys()
                ):
                    if selected_doc in list(json_db[user][domain].keys()):
                        doc_sections = json_db[user][domain][selected_doc]

            if len(list(doc_sections.keys())) > 0:
                # Duplicate Sections
                dup_sections = [
                    {
                        "dup_section": doc_sections[section_key]["section"],
                        "content": doc_sections[section_key]["content"],
                    }
                    for section_key in doc_sections.keys()
                    if re.search(
                        r"0.[a-z]$", doc_sections[section_key]["section"].split(" ")[0]
                    )
                ]
                if len(dup_sections) > 0:
                    for dup_sec in dup_sections:
                        cur_sec = dup_sec["dup_section"].split(" ")[0]
                        sec_ls = cur_sec.split(".")
                        parent_sec = ".".join(sec_ls[0 : len(sec_ls) - 2])
                        for section_key in doc_sections.keys():
                            section_title = doc_sections[section_key]["section"]
                            section_ = section_title.split(" ")[0]
                            if section_ == parent_sec:
                                sections_json_db[selected_doc][section_title][
                                    "content"
                                ] += dup_sec["content"]
                final_sel_sections = []
                for section_key in doc_sections.keys():
                    section_title = doc_sections[section_key]["section"]
                    for sel_section in selected_sections:
                        # Add Parent Section if Duplicate is selected here!
                        if section_title == sel_section and (
                            not re.search(r"0.[a-z]$", section_title.split(" ")[0])
                        ):
                            if (
                                sections_json_db != None
                                and selected_doc in sections_json_db.keys()
                                and section_title
                                in sections_json_db[selected_doc].keys()
                            ):
                                doc_content.append(
                                    sections_json_db[selected_doc][section_title][
                                        "content"
                                    ]
                                )
                            else:
                                doc_content.append(doc_sections[section_key]["content"])
                            final_sel_sections.append(sel_section)

        return {
            "status": "Received!",
            "sel_sec_content": doc_content,
            "sel_sections": final_sel_sections,
        }
    except Exception as exp:
        return {"status": f"Error occured as {exp}"}


@application.route("/get_uploaded_dh_section_content", methods=["GET", "POST"])
def get_uploaded_dh_section_content():
    try:
        prop = request.get_json()
        section_ = prop["section"]
        filename = prop["filename"]
        domain = prop["domain"]
        user = prop["user"]
        doc_content = ""
        is_file_exists = True
        if is_file_exists:
            section_json = {}
            sections_db = sections_mappings.find_one(
                format_data(filename, None, user)["filter"]
            )
            if sections_db != None:
                section_json = dict(sections_db)
                section_json = bson_to_json(
                    "sections_mappings", section_json, filename, None, user
                )
                section_json = {user: section_json}

            if user in section_json.keys() and filename in section_json[user].keys():
                sections = section_json[user][filename]
                if section_ in sections.keys():
                    content = sections[section_]["content"]
                    return {
                        "status": "Received!",
                        "content": content,
                        "section": section_,
                    }

            uploaded_sections_db = {}
            history_db = download_history_mappings.find_one(
                format_data(filename, domain, user)["filter"]
            )
            if history_db != None:
                uploaded_sections_db = dict(history_db)
                uploaded_sections_db = bson_to_json(
                    "download_history_mappings",
                    uploaded_sections_db,
                    filename,
                    domain,
                    user,
                )

            if (
                user in uploaded_sections_db.keys()
                and domain in uploaded_sections_db[user].keys()
                and filename in uploaded_sections_db[user][domain].keys()
            ):
                doc_sections = uploaded_sections_db[user][domain][filename]
                if len(list(doc_sections.keys())) > 0:
                    for section_key in doc_sections.keys():
                        section_title = doc_sections[section_key]["section"]
                        if section_title == section_:
                            doc_content = doc_sections[section_key]["content"]

        return {"status": "Received!", "content": doc_content}
    except Exception as exp:
        return {"status": f"Error occured as {exp}"}


@application.route("/set_uploaded_dh_section_content", methods=["GET", "POST"])
def set_uploaded_dh_section_content():
    try:
        prop = request.get_json()
        section_ = prop["section"]
        filename = prop["filename"]
        content = prop["content"]
        user = prop["user"]

        is_file_exists = True
        if is_file_exists:
            section_json = {}
            sections_db = sections_mappings.find_one(
                format_data(filename, None, user)["filter"]
            )
            if sections_db != None:
                section_json = dict(sections_db)
                section_json = bson_to_json(
                    "sections_mappings", section_json, filename, None, user
                )

            if user in section_json.keys() and filename in section_json[user].keys():
                sections = section_json[user][filename]
                if section_ in sections.keys():
                    sections[section_]["content"] = content
                else:
                    sections[section_] = {"content": content}
            else:
                section_json[filename] = {section_: {"content": content}}

            sections_bson = json_to_bson(
                {"sections": section_json[filename]},
                format_data(filename, None, user)["filter"],
            )
            if bool(sections_db):
                sections_mappings.update_one(
                    format_data(filename, None, user)["filter"],
                    {"$set": {"sections": section_json[filename]}},
                )
            else:
                sections_mappings.insert_many(sections_bson)

        return {"status": "Received!"}
    except Exception as exp:
        return {"status": f"Error occured as {exp}"}


@application.route("/download_with_json", methods=["GET", "POST"])
def download_with_json():
    try:
        prop = request.get_json()
        # domain = prop["domain"]
        file_id = prop["file_id"].split("-")
        ascii_dec_domain = file_id[len(file_id) - 1].encode("ascii")
        base64_dec_domain = base64.b64decode(ascii_dec_domain)
        domain = base64_dec_domain.decode("ascii")
        filename = "-".join(file_id[0 : len(file_id) - 1])
        key_map = prop["key_map"]
        is_crm = prop["is_crm"]
        crmid = prop["crmid"]
        # Custom vars mappings
        custom_vars_dict = {}
        is_file_exists = True
        if is_file_exists:
            custom_vars_db = custom_vars_mappings.find_one(
                format_data(filename, domain)["filter"]
            )
            if custom_vars_db != None:
                json_db = dict(custom_vars_db)
                json_db = bson_to_json(
                    "custom_vars_mappings", json_db, filename, domain
                )

            if json_db != None:
                if domain in json_db.keys() and filename in json_db[domain].keys():
                    custom_vars_dict = json_db[domain][filename]
        custom_vars = custom_vars_dict.keys()
        if len(custom_vars) == 0:
            custom_vars = [
                custom_var
                for custom_var in prop.keys()
                if custom_var not in ["file_id", "key_map", "crmid", "is_crm"]
            ]
        # Get BBT Data
        if is_crm:
            bbt_data = get_bbt_data(crmid)
            if len(bbt_data.keys()) > 0:
                pass
            else:
                raise Exception("Invalid CRM-ID!")
        else:
            bbt_data = {
                "customer": "Test customer",
                "market": "Test market",
                "opportunityname": "Test opportunity name",
                "country": "Test country",
                "region": "Test region",
                "subregion": "Test subregion",
            }
        bbt_data["crmid"] = crmid

        if key_map != None:
            data_key_map = [
                f"{cur_['parameter']}_{cur_['dependency']}_{cur_['value']}"
                for cur_ in key_map
            ]

            # Call Download Doc API for respective contents and sections
            payload = {"domain": domain, "filename": filename, "key_map": data_key_map}
            res = requests.post(
                "http://localhost:5000/doc-assist/download_doc",
                json=payload,
                headers={"Content-Type": "application/json"},
            )
            res = res.json()
            down_contents = res["contents"]
            down_sections = res["sections"]
            # Add Empty Table for Array Custom Vars
            down_contents_ = []
            for content in down_contents:
                cur_content = content
                for cvar in custom_vars:
                    cvar_pat = re.sub("\[", "\[", cvar)
                    cvar_pat = re.sub("\]", "\]", cvar_pat)
                    if len(re.findall(cvar_pat, cur_content)) > 0:
                        if (
                            type(custom_vars_dict[cvar]) != str
                            and custom_vars_dict[cvar]["crm_bg_option"] == "array"
                            and cvar in prop.keys()
                        ):
                            table_ = ""
                            for row in prop[cvar]:
                                table_ += "<tr>"
                                for _ in row:
                                    table_ += f"<td>{cvar}</td>"
                                table_ += "</tr>"
                            cur_content = cur_content.replace(
                                cvar,
                                f"<table>{table_}</table>",
                            )
                        elif (
                            type(custom_vars_dict[cvar]) != str
                            and custom_vars_dict[cvar]["crm_bg_option"] == "img"
                            and cvar in prop.keys()
                            and isBase64(prop[cvar].split("base64,")[1])
                        ):
                            cur_content = cur_content.replace(
                                cvar,
                                f"<img src='{prop[cvar]}'/>",
                            )
                down_contents_.append(cur_content)
            down_sections = format_sections(down_sections)
            final_doc = format_html(down_sections, down_contents_)

            conv_html = final_doc
            conv_to = "docx"
            parser = DocHTMLParser()
            parser.set_document()
            parser.set_fig_count()
            parser.set_table_count()
            parser.set_prev_start_tag()
            parser.set_table_tags()
            parser.feed(conv_html)
            document = parser.get_document()
            for para in document.paragraphs:
                if len(re.findall("Figure [0-9]*:", para.text)) > 0:
                    fig_tab_caption(para, True)
                elif len(re.findall("Table [0-9]*:", para.text)) > 0:
                    fig_tab_caption(para)
            document.save("doc_assist_be/services/converted_doc.docx")
            document = Document("doc_assist_be/services/converted_doc.docx")
            for para in document.paragraphs:
                print(para)
                if para.text == "%TOC%":
                    para.text = ""
                    run = para.add_run()
                    add_list_of_table(run)
                if para.text == "%TOF%":
                    para.text = ""
                    run = para.add_run()
                    add_list_of_fig_tab(run, True)
                if para.text == "%TOT%":
                    para.text = ""
                    run = para.add_run()
                    add_list_of_fig_tab(run)
                else:
                    update_custom_vars(
                        bbt_data,
                        custom_vars_dict,
                        custom_vars,
                        para,
                        prop,
                        -1,
                        -1,
                    )

            # table = document.tables[0]
            for table in document.tables:
                for row_idx, each_row in enumerate(table.rows):
                    for ecell_idx, each_cell in enumerate(each_row.cells):
                        start_idx = 0
                        if (
                            len(each_cell.paragraphs[0].text) == 0
                            and len(each_cell.paragraphs) > 1
                        ):
                            delete_paragraph(each_cell.paragraphs[0])
                            # start_idx = 1
                        if row_idx == 0:
                            if len(each_cell.paragraphs) > 0:
                                for para in each_cell.paragraphs[start_idx:]:
                                    update_custom_vars(
                                        bbt_data,
                                        custom_vars_dict,
                                        custom_vars,
                                        para,
                                        prop,
                                        row_idx,
                                        ecell_idx,
                                    )
                                    format_table(each_cell, para, True)
                            set_table_header_bg_color(each_cell, "#005AFF")
                        else:
                            if len(each_cell.paragraphs) > 0:
                                for para in each_cell.paragraphs[start_idx:]:
                                    update_custom_vars(
                                        bbt_data,
                                        custom_vars_dict,
                                        custom_vars,
                                        para,
                                        prop,
                                        row_idx,
                                        ecell_idx,
                                    )
                            #         format_table(each_cell, para, False)
                            set_table_header_bg_color(each_cell, "#CCCCCC")
                        set_cell_border(
                            each_cell,
                            top={"sz": 12, "color": "#FFFFFF", "val": "single"},
                            bottom={"sz": 12, "color": "#FFFFFF", "val": "single"},
                            start={"sz": 12, "color": "#FFFFFF", "val": "single"},
                            end={"sz": 12, "color": "#FFFFFF", "val": "single"},
                        )
            document.save("doc_assist_be/services/converted_doc.docx")
            # pythoncom.CoInitialize()
            # word = comtypes.client.CreateObject(
            #     "Word.Application"
            # )  # opens the word application
            # doc = word.Documents.Open(filepath)  # opens the specified file
            # res = doc.Fields.Update()  # updates field, returns 0 if successful
            # print(res)
            # doc.Save()
            # doc.Close()
            # word.Quit()
            if conv_to == "pdf":
                try:
                    convert(filepath)
                finally:
                    return send_file("doc_assist_be/services/converted_doc.pdf")
            return send_file("doc_assist_be/services/converted_doc.docx")
        return {"status": "Received!"}
    except Exception as exp:
        return {"status": f"Error occured as {exp}"}


@application.route("/validate_crm", methods=["GET", "POST"])
def validate_crm():
    try:
        prop = request.get_json()
        domain = prop["domain"]
        filename = prop["filename"]
        is_crm = prop["is_crm"]
        crmid = prop["crmid"]

        # Get BBT Data
        if is_crm:
            bbt_data = get_bbt_data(crmid)
            if len(bbt_data.keys()) > 0:
                pass
            else:
                raise Exception("Invalid CRM-ID!")
        else:
            bbt_data = {
                "customer": "Test customer",
                "market": "Test market",
                "opportunityname": "Test opportunity name",
                "country": "Test country",
                "region": "Test region",
                "subregion": "Test subregion",
            }
        bbt_data["crmid"] = crmid

        # Custom vars mappings
        custom_vars_dict = {}
        custom_vars = []
        custom_vars_payload = {}
        is_file_exists = os.path.isfile(custom_vars_mappings)
        if is_file_exists:
            with open(custom_vars_mappings, "r") as outfile:
                json_db = json.load(outfile)

            if json_db != None:
                if domain in json_db.keys() and filename in json_db[domain].keys():
                    custom_vars_dict = json_db[domain][filename]
                    custom_vars = custom_vars_dict.keys()

                    for cvar in custom_vars:
                        if custom_vars_dict[cvar]["crm_bg"] == "crm":
                            option = custom_vars_dict[cvar]["crm_bg_option"]
                            if option in bbt_data.keys():
                                custom_vars_payload[cvar] = bbt_data[option]

        return {"status": "Received!", "bbt_mapping": custom_vars_payload}
    except Exception as exp:
        return {"status": f"Error occured as {exp}"}


@application.route("/set_update_file_cache", methods=["GET", "POST"])
def set_update_file_cache():
    try:
        prop = request.get_json()
        domain = prop["domain"]
        selected_doc = prop["selectedDoc"]
        user = prop["user"]
        # user = prop["user"]
        is_file_exists = True
        doc_sections = {}
        if is_file_exists:
            json_db = {}
            master_db = master_mappings.find_one(
                format_data(selected_doc, domain)["filter"]
            )
            if master_db != None:
                json_db = dict(master_db)
                json_db = bson_to_json("master_mappings", json_db, selected_doc, domain)

            if json_db != None and domain in json_db.keys():
                if selected_doc in list(json_db[domain].keys()):
                    doc_sections = json_db[domain][selected_doc]

        formatted_sections = {}
        version_hierarchy = {}

        if len(doc_sections.keys()) > 0:
            for doc_sec in doc_sections.keys():
                sec_title_idx = doc_sections[doc_sec]["section"].split(" ")

                version = sec_title_idx[0]
                node = {"children": []}
                sub_vers = version.split(".")
                node_ver = ".".join(sub_vers)
                if len(sub_vers) > 1:
                    parent = ".".join(sub_vers[0 : len(sub_vers) - 1])
                    if parent not in version_hierarchy.keys():
                        version_hierarchy[parent] = {"children": []}
                    version_hierarchy[parent]["children"].append(node_ver)
                    version_hierarchy[node_ver] = node
                    # print(f"Parent: {parent} - Child: {version}")
                else:
                    version_hierarchy[node_ver] = node

                param_dep_val = doc_sec.split("_")
                formatted_sections[sec_title_idx[0]] = {
                    "parameter": param_dep_val[0],
                    "dependency": param_dep_val[1],
                    "value": param_dep_val[2:],
                    "content": doc_sections[doc_sec]["content"],
                    "title": " ".join(sec_title_idx[1:]),
                    "id": f"#_Toc_{sec_title_idx[0]}",
                }

        # Update Cache
        if is_file_exists:
            sections_db = {}
            section_db = sections_mappings.find_one(
                format_data(selected_doc, None, user)["filter"]
            )
            if section_db != None:
                sections_db = dict(section_db)
                sections_db = bson_to_json(
                    "sections_mappings", sections_db, selected_doc, None, user
                )

            if sections_db != None:
                sections_db[selected_doc] = formatted_sections

            sections_bson = json_to_bson(
                {"sections": sections_db[selected_doc]},
                format_data(selected_doc, None, user)["filter"],
            )
            if bool(sections_db):
                sections_mappings.update_one(
                    format_data(selected_doc, None, user)["filter"],
                    {"$set": {"sections": sections_db[selected_doc]}},
                )
            else:
                sections_mappings.insert_many(sections_bson)

        # Update Heirarchy Mapping
        if is_file_exists:
            heirarchy_db = {}
            version_h_db = heirarchy_mappings.find_one(
                format_data(selected_doc, None, user)["filter"]
            )
            if version_h_db != None:
                heirarchy_db = dict(version_h_db)
                heirarchy_db = bson_to_json(
                    "heirarchy_mappings", heirarchy_db, selected_doc, None, user
                )
                # if heirarchy_db != None:
                #     heirarchy_db[selected_doc] = version_hierarchy
            versions_bson = json_to_bson(
                {"version_heirarchy": version_hierarchy},
                format_data(selected_doc, None, user)["filter"],
            )
            if bool(version_h_db):
                heirarchy_mappings.update_one(
                    format_data(selected_doc, None, user)["filter"],
                    {"$set": {"version_heirarchy": heirarchy_db[selected_doc]}},
                )
            else:
                heirarchy_mappings.insert_many(versions_bson)

        return {"status": "Received!"}

    except Exception as exp:
        return {"status": f"Error occured as {exp}"}
